"""Multi-Format Dokument-Parser für User-Dokumente.

Unterstützte Formate:
- PDF (via pymupdf/fitz)
- DOCX (via python-docx)
- LaTeX (.tex) — Regex-basiertes Stripping von LaTeX-Kommandos
- TXT/MD — Plain-Text
- Bilder (PNG, JPG, JPEG, TIFF) — als vision_pending markiert (OCR in AP-05)
"""

import logging
import re
from pathlib import Path

from pydantic import BaseModel

logger = logging.getLogger(__name__)

# ─── Datenmodell ─────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS: set[str] = {
    ".pdf",
    ".docx",
    ".tex",
    ".txt",
    ".md",
    ".png",
    ".jpg",
    ".jpeg",
    ".tiff",
    ".tif",
}

IMAGE_EXTENSIONS: set[str] = {".png", ".jpg", ".jpeg", ".tiff", ".tif"}

DOC_TYPE_PATTERNS: dict[str, list[str]] = {
    "cv": [
        "cv",
        "lebenslauf",
        "resume",
        "curriculum",
        "vita",
    ],
    "zeugnis": [
        "zeugnis",
        "arbeitszeugnis",
        "certificate",
        "reference",
        "zwischenzeugnis",
        "endzeugnis",
        "dienstzeugnis",
    ],
    "zertifikat": [
        "zertifikat",
        "cert",
        "nachweis",
        "bescheinigung",
        "weiterbildung",
        "schulung",
    ],
    "projekt": [
        "projekt",
        "project",
        "portfolio",
    ],
}


class ParsedDocument(BaseModel):
    """Ein geparster Dokument-Inhalt."""

    path: str
    filename: str
    doc_type: str  # 'cv'|'zeugnis'|'zertifikat'|'projekt'|'sonstiges'
    text: str
    page_count: int | None = None
    parse_method: str  # 'pymupdf'|'python-docx'|'latex_strip'|'plain'|'vision_pending'


# ─── Dokument-Parser ────────────────────────────────────────────────────────


class DocumentParser:
    """Parst Dokumente verschiedener Formate zu Text."""

    async def parse_file(self, path: Path) -> ParsedDocument:
        """Dispatche anhand der Dateiendung zum passenden Parser.

        Args:
            path: Pfad zur Datei

        Returns:
            ParsedDocument mit extrahiertem Text

        Raises:
            FileNotFoundError: Datei existiert nicht
            ValueError: Nicht unterstütztes Dateiformat
        """
        if not path.exists():
            raise FileNotFoundError(f"Datei nicht gefunden: {path}")

        suffix = path.suffix.lower()

        if suffix not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Nicht unterstütztes Format: {suffix}. "
                f"Unterstützt: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        filename = path.name
        doc_type = self._classify_doc_type(filename)

        # Bilder → vision_pending (OCR/Vision in AP-05)
        if suffix in IMAGE_EXTENSIONS:
            return ParsedDocument(
                path=str(path),
                filename=filename,
                doc_type=doc_type,
                text="",
                page_count=None,
                parse_method="vision_pending",
            )

        # Text-basierte Formate
        if suffix == ".pdf":
            text, page_count = self._parse_pdf(path)
            return ParsedDocument(
                path=str(path),
                filename=filename,
                doc_type=doc_type,
                text=text,
                page_count=page_count,
                parse_method="pymupdf",
            )

        if suffix == ".docx":
            text = self._parse_docx(path)
            return ParsedDocument(
                path=str(path),
                filename=filename,
                doc_type=doc_type,
                text=text,
                page_count=None,
                parse_method="python-docx",
            )

        if suffix == ".tex":
            text = self._parse_latex(path)
            return ParsedDocument(
                path=str(path),
                filename=filename,
                doc_type=doc_type,
                text=text,
                page_count=None,
                parse_method="latex_strip",
            )

        # .txt, .md → Plain-Text
        text = self._parse_txt(path)
        return ParsedDocument(
            path=str(path),
            filename=filename,
            doc_type=doc_type,
            text=text,
            page_count=None,
            parse_method="plain",
        )

    async def scan_folder(self, folder: Path) -> list[ParsedDocument]:
        """Scanne einen Ordner rekursiv nach unterstützten Dateien.

        Args:
            folder: Pfad zum Ordner

        Returns:
            Liste von ParsedDocument für alle unterstützten Dateien

        Raises:
            FileNotFoundError: Ordner existiert nicht
        """
        if not folder.exists():
            raise FileNotFoundError(f"Ordner nicht gefunden: {folder}")
        if not folder.is_dir():
            raise ValueError(f"Kein Ordner: {folder}")

        documents: list[ParsedDocument] = []
        for file_path in sorted(folder.rglob("*")):
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                logger.debug("Überspringe nicht unterstützte Datei: %s", file_path.name)
                continue
            # Versteckte Dateien und __pycache__ überspringen
            if any(part.startswith(".") or part == "__pycache__" for part in file_path.parts):
                continue

            try:
                doc = await self.parse_file(file_path)
                documents.append(doc)
                logger.info(
                    "Geparst: %s → %s (%s, %d Zeichen)",
                    file_path.name,
                    doc.doc_type,
                    doc.parse_method,
                    len(doc.text),
                )
            except Exception:
                logger.exception("Fehler beim Parsen von %s", file_path)

        logger.info(
            "Ordner-Scan abgeschlossen: %d Dokumente in %s",
            len(documents),
            folder,
        )
        return documents

    def _parse_pdf(self, path: Path) -> tuple[str, int]:
        """Extrahiere Text aus einer PDF-Datei via pymupdf.

        Returns:
            Tuple von (text, page_count)
        """
        import fitz  # pymupdf

        doc = fitz.open(str(path))
        pages: list[str] = []
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                pages.append(text.strip())
        page_count = len(doc)
        doc.close()

        full_text = "\n\n".join(pages)
        return full_text, page_count

    def _parse_docx(self, path: Path) -> str:
        """Extrahiere Text aus einer DOCX-Datei via python-docx."""
        from docx import Document

        doc = Document(str(path))
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Tabellen ebenfalls extrahieren
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs)

    def _parse_latex(self, path: Path) -> str:
        r"""Extrahiere Text aus einer LaTeX-Datei via Regex-Stripping.

        Entfernt LaTeX-Kommandos und gibt den Rohtext zurück.
        Kein vollständiger LaTeX-Parser — reicht für Profil-Extraktion.
        """
        raw = path.read_text(encoding="utf-8", errors="replace")

        # Kommentare entfernen (aber nicht \%)
        text = re.sub(r"(?<!\\)%.*$", "", raw, flags=re.MULTILINE)

        # Häufige Umgebungen entfernen (begin/end)
        text = re.sub(r"\\begin\{[^}]+\}", "", text)
        text = re.sub(r"\\end\{[^}]+\}", "", text)

        # LaTeX-Kommandos mit einem Argument: \command{text} → text
        text = re.sub(
            r"\\(?:textbf|textit|emph|underline|section|subsection"
            r"|subsubsection|paragraph|href)\{([^}]*)\}",
            r"\1",
            text,
        )

        # LaTeX-Kommandos mit zwei Argumenten: \command{a}{b} → b
        text = re.sub(r"\\href\{[^}]*\}\{([^}]*)\}", r"\1", text)

        # Verbleibende Kommandos ohne Argumente entfernen
        text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", text)

        # Geschweifte Klammern entfernen
        text = re.sub(r"[{}]", "", text)

        # Spezialzeichen normalisieren
        text = text.replace("~", " ")
        text = text.replace("\\\\", "\n")
        text = text.replace("\\&", "&")
        text = text.replace("\\%", "%")
        text = text.replace("\\_", "_")

        # Mehrfache Leerzeilen auf maximal zwei reduzieren
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Führende/trailing Whitespace pro Zeile entfernen
        lines = [line.strip() for line in text.splitlines()]
        text = "\n".join(line for line in lines if line)

        return text

    def _parse_txt(self, path: Path) -> str:
        """Lese Plain-Text-Datei."""
        return path.read_text(encoding="utf-8", errors="replace").strip()

    def _classify_doc_type(self, filename: str) -> str:
        """Klassifiziere den Dokumenttyp anhand des Dateinamens.

        Prüft den Dateinamen (ohne Extension, case-insensitive) gegen
        bekannte Muster. Bei mehreren Treffern gewinnt die spezifischere
        Kategorie (cv > zeugnis > zertifikat > projekt).
        """
        name_lower = filename.lower().rsplit(".", 1)[0]  # ohne Extension
        # Unterstriche und Bindestriche als Wortgrenzen behandeln
        name_normalized = re.sub(r"[-_]", " ", name_lower)

        for doc_type, patterns in DOC_TYPE_PATTERNS.items():
            for pattern in patterns:
                if pattern in name_normalized:
                    return doc_type

        return "sonstiges"
